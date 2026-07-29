#!/usr/bin/env python3
"""生成课堂小结 .docx 文档。

用法:
  python create_docx.py --title "火箭2班数学小结7.14" \
      --content '["内容1", "内容2", "内容3"]' \
      --requirements '["要求1", "要求2", "要求3"]' \
      --output "C:\\...\\课堂小结0714.docx"
"""

import argparse
import json
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_run_font(run, font_name="等线", font_name_east_asian="等线", size_pt=11):
    """设置 run 的中英文字体和字号。"""
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name_east_asian)


def create_docx(title, content_items, requirement_items, output_path):
    """生成课堂小结文档。

    Args:
        title: 标题行，如 "火箭2班数学小结7.14"
        content_items: 授课内容列表（3条）
        requirement_items: 学习要求列表（3条）
        output_path: 输出 .docx 文件路径
    """
    doc = Document()

    # --- 设置默认字体 ---
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "等线"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "等线")

    # --- 页面设置（与模板一致：A4, 标准边距）---
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.175)
        section.right_margin = Cm(3.175)

    # --- 标题行 ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_title.add_run(title)
    set_run_font(run, size_pt=11)

    # --- 空行 ---
    doc.add_paragraph()

    # --- 授课内容 ---
    p_section1 = doc.add_paragraph()
    run = p_section1.add_run("1️⃣ 授课内容")
    set_run_font(run, size_pt=11)

    circled_numbers = ["①", "②", "③"]
    for i, item in enumerate(content_items):
        p = doc.add_paragraph()
        run = p.add_run(f"{circled_numbers[i]} {item}")
        set_run_font(run, size_pt=11)

    # --- 空行 ---
    doc.add_paragraph()

    # --- 学习要求 ---
    p_section2 = doc.add_paragraph()
    run = p_section2.add_run("2️⃣ 学习要求")
    set_run_font(run, size_pt=11)

    for i, item in enumerate(requirement_items):
        p = doc.add_paragraph()
        run = p.add_run(f"{circled_numbers[i]} {item}")
        set_run_font(run, size_pt=11)

    # --- 末尾空行 ---
    doc.add_paragraph()

    # --- 保存 ---
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"课堂小结已保存至: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="生成课堂小结 .docx 文档")
    parser.add_argument("--title", required=True, help="标题行")
    parser.add_argument("--content", required=True, help="授课内容 JSON 数组（3条）")
    parser.add_argument("--requirements", required=True, help="学习要求 JSON 数组（3条）")
    parser.add_argument("--output", required=True, help="输出文件路径")
    args = parser.parse_args()

    content_items = json.loads(args.content)
    requirement_items = json.loads(args.requirements)

    if len(content_items) != 3:
        print(f"[警告] 授课内容应为 3 条，当前为 {len(content_items)} 条")
    if len(requirement_items) != 3:
        print(f"[警告] 学习要求应为 3 条，当前为 {len(requirement_items)} 条")

    create_docx(args.title, content_items, requirement_items, args.output)


if __name__ == "__main__":
    main()
